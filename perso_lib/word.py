import os
import six
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
try:
    from html import escape
except ImportError:
    # cgi.escape is deprecated in python 3.7
    from cgi import escape

class Docx:
    NEWLINE_XML = '</w:t><w:br/><w:t xml:space="preserve">'
    NEWPARAGRAPH_XML = '</w:t></w:r></w:p><w:p><w:r><w:t xml:space="preserve">'
    TAB_XML = '</w:t></w:r><w:r><w:tab/></w:r><w:r><w:t xml:space="preserve">'
    PAGE_BREAK = '</w:t><w:br w:type="page"/><w:t xml:space="preserve">'
    def __init__(self,doc_name):
        if doc_name:
            self.doc = Document(doc_name)
        else:
            self.doc = Document()
        self.doc_name = doc_name

    def add_heading(self,title_level,title):
        '''
        添加标题
        title_level 表示添加"标题xx",itle_level为1,2,3...
        title 标题名称
        '''
        self.doc.add_heading(title, level=title_level)

    def add_paragraph(self,content):
        return self.doc.add_paragraph(content)

    def set_paragraph_property(self,paragraph,blod,font_size,font_color,bk_color):
        pass

    def add_table(self,row=1,col=1,style=None):
        table = self.doc.add_table(rows=row,cols=col, style='Table Grid')
        table.autofit = False
        return table

    def get_cell(self,table,row,col):
        '''
        获取表格指定行和列的单元格
        table为_Table对象，可通过add_table函数获取
        row 为cell所在的行，索引从0,1,2...表示第1,2,3...行
        colo为cell所在的列，索引从0,1,2...表示第1,2,3...列
        '''
        return table.rows[row].cells[col]
        # return table.cell(row,col)

    def merge_cell(self,cells):
        merged_cell = cells[0]
        if len(cells) < 2:
            return merged_cell
        else:
            text = ''
            for i in range(1,len(cells)):
                merged_cell = merged_cell.merge(cells[i])
                if merged_cell.paragraphs and len(merged_cell.paragraphs) > 0:
                    text = merged_cell.text.replace('\n','')
            merged_cell.text = text
        return merged_cell

    
    def set_cell_background(self,cell,bk_color=None):
        '''
        设置单元格背景颜色
        cell:       _Cell对象,可通过get_cell函数获取
        bk_color:   RGB字符串，例如"AABBCC"
        '''
        xml = r'<w:shd {0} w:fill="{1}"/>'.format(nsdecls('w'),bk_color)
        print(xml)
        shading_elm_1 = parse_xml(xml)
        cell._tc.get_or_add_tcPr().append(shading_elm_1)

    def set_cell_text(self,cell,text):
        '''
        设置单元格文本
        cell    _Cell对象,可通过get_cell函数获取
        text    需要添加的文本内容，若cell中已经有文本，则被替换掉
        '''
        paragraph = None
        if cell.paragraphs and len(cell.paragraphs) > 0:
            paragraph = cell.paragraphs[-1] #取最后一个索引的段落
            paragraph.text = text
        else:
            paragraph = cell.add_paragraph(text)
        paragraph.first_line_indent = Inches(-10)
        paragraph.paragraph_format.first_line_indent = Inches(0)
        

        # p.aligment = WD_ALIGN_PARAGRAPH.LEFT
        # p.left_indent = Inches(0)
        # paragraph_format = p.paragraph_format
        # paragraph_format.left_indent = Inches(0)
        # cell.text = text

    def set_cell(self,cell,text,
                        style=None,
                        color=None,
                        highlight=None,
                        size=None,
                        subscript=None,
                        superscript=None,
                        bold=False,
                        italic=False,
                        underline=False,
                        strike=False,
                        font=None,
                        url_id=None):
        # If not a string : cast to string (ex: int, dict etc...)
        if not isinstance(text, (six.text_type, six.binary_type)):
            text = six.text_type(text)
        if not isinstance(text, six.text_type):
            text = text.decode('utf-8',errors='ignore')
        text = (escape(text)
                .replace('\n', self.NEWLINE_XML)
                .replace('\a', self.NEWPARAGRAPH_XML)
                .replace('\t', self.TAB_XML)
                .replace('\f', self.PAGE_BREAK))

        prop = u''

        if style:
            prop += u'<w:rStyle w:val="%s"/>' % style
        if color:
            if color[0] == '#':
                color = color[1:]
            prop += u'<w:color w:val="%s"/>' % color
        if highlight:
            if highlight[0] == '#':
                highlight = highlight[1:]
            prop += u'<w:highlight w:val="%s"/>' % highlight
        if size:
            prop += u'<w:sz w:val="%s"/>' % size
            prop += u'<w:szCs w:val="%s"/>' % size
        if subscript:
            prop += u'<w:vertAlign w:val="subscript"/>'
        if superscript:
            prop += u'<w:vertAlign w:val="superscript"/>'
        if bold:
            prop += u'<w:b/>'
        if italic:
            prop += u'<w:i/>'
        if underline:
            if underline not in ['single','double']:
                underline = 'single'
            prop += u'<w:u w:val="%s"/>' % underline
        if strike:
            prop += u'<w:strike/>'
        if font:
            prop += ( u'<w:rFonts w:ascii="{font}" w:hAnsi="{font}" w:cs="{font}"/>'
                      .format(font=font) )


        xml = u'<w:r>'
        if prop:
            xml += u'<w:rPr>%s</w:rPr>' % prop
        xml += u'<w:t xml:space="preserve">%s</w:t></w:r>' % text
        if url_id:
            xml = ( u'<w:hyperlink r:id="%s" w:tgtFrame="_blank">%s</w:hyperlink>'
                    % (url_id, xml) )
        print(xml)
        tcPr = parse_xml(xml)
        cell._tc.get_or_add_tcPr().append(tcPr)


    def close(self):
        self.doc.save(self.doc_name)
    
    def save_as(self,doc_name):
        self.doc.save(doc_name)


if __name__ == '__main__':
    doc_path ='D:\\DP需求.docx'
    document = Docx(doc_path)
    document.add_heading(2,'Visa 分组')
    document.add_heading(3,'DGI0101')
    new_table = document.add_table(3,4)
    for col in range(4):
        cell = document.get_cell(new_table,0,col)
        document.set_cell_background(cell,'FFCA00')
        # document.set_cell(cell,'abc',color='FFCA00')
    for row in range(1,3):
        for col in range(4):
            cell = document.get_cell(new_table,row,col)
            document.set_cell_text(cell,'AA')
    new_row = new_table.add_row()
    for cell in new_row.cells:
        document.set_cell_text(cell,'BB')

    cell1 = document.get_cell(new_table,1,1)
    cell2 = document.get_cell(new_table,1,2)
    merged_cell = document.merge_cell([cell1,cell2])
    document.save_as('D:\\demo.docx')